from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_template(name, style, sections):
    doc = Document()
    
    # Add header
    header = doc.sections[0].header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_paragraph.add_run(f"{name} Template").bold = True
    
    # Add sections
    for section in sections:
        # Add section title
        title = doc.add_heading(section['title'], level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add content
        if 'content' in section:
            for item in section['content']:
                p = doc.add_paragraph()
                run = p.add_run(item)
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add skills if present
        if 'skills' in section:
            doc.add_paragraph("Skills:")
            for skill in section['skills']:
                p = doc.add_paragraph()
                run = p.add_run(f"• {skill}")
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add divider
        doc.add_paragraph("—" * 80)
    
    # Save the document
    doc.save(f'{name.lower().replace(" ", "-")}.docx')

# Modern Professional Template
modern_sections = [
    {
        'title': 'Contact Information',
        'content': [
            'John Doe',
            '123 Main Street',
            'City, State 12345',
            'Phone: (555) 123-4567',
            'Email: john.doe@example.com',
            'LinkedIn: linkedin.com/in/johndoe'
        ]
    },
    {
        'title': 'Professional Summary',
        'content': [
            'Results-driven professional with 5+ years of experience in project management and team leadership. Proven track record of delivering successful projects on time and under budget.'
        ]
    },
    {
        'title': 'Core Competencies',
        'skills': [
            'Project Management', 'Team Leadership', 'Budget Management',
            'Risk Assessment', 'Stakeholder Communication', 'Process Improvement'
        ]
    },
    {
        'title': 'Professional Experience',
        'content': [
            'Senior Project Manager',
            'ABC Corporation',
            'City, State',
            '2018 - Present',
            '• Led cross-functional teams of 10+ members',
            '• Managed project budgets of $1M+',
            '• Implemented process improvements that increased efficiency by 20%'
        ]
    }
]

# Technical Engineering Template
tech_sections = [
    {
        'title': 'Technical Skills',
        'skills': [
            'Python', 'Java', 'C++', 'SQL', 'AWS', 'Docker',
            'Kubernetes', 'CI/CD', 'REST APIs', 'Microservices'
        ]
    },
    {
        'title': 'Professional Experience',
        'content': [
            'Senior Software Engineer',
            'TechCorp',
            'City, State',
            '2019 - Present',
            '• Led development of microservices architecture',
            '• Implemented CI/CD pipeline using Jenkins',
            '• Optimized database queries for 50% performance improvement'
        ]
    },
    {
        'title': 'Education',
        'content': [
            'Master of Computer Science',
            'University of Technology',
            'Graduated: 2018',
            'GPA: 3.9/4.0'
        ]
    },
    {
        'title': 'Projects',
        'content': [
            '• Developed real-time analytics platform using Python and Kafka',
            '• Created automated testing framework for REST APIs',
            '• Implemented container orchestration system with Kubernetes'
        ]
    }
]

# Creative Design Template
creative_sections = [
    {
        'title': 'Design Skills',
        'skills': [
            'Adobe Creative Suite', 'UI/UX Design', 'Brand Identity',
            'Illustration', 'Typography', 'Motion Graphics',
            'Print Design', 'Digital Marketing'
        ]
    },
    {
        'title': 'Portfolio Highlights',
        'content': [
            '• Created award-winning brand identity for startup',
            '• Designed responsive websites with modern UX',
            '• Developed comprehensive marketing materials'
        ]
    },
    {
        'title': 'Education & Certifications',
        'content': [
            'Bachelor of Fine Arts',
            'Design Academy',
            'Graduated: 2017',
            'Certified Adobe Creative Suite Expert'
        ]
    },
    {
        'title': 'Professional Experience',
        'content': [
            'Senior Graphic Designer',
            'Creative Agency',
            'City, State',
            '2018 - Present',
            '• Led design team of 5 designers',
            '• Created successful marketing campaigns',
            '• Developed brand guidelines for major clients'
        ]
    }
]

# Create templates
create_template("Modern Professional", "modern", modern_sections)
create_template("Technical Engineering", "technical", tech_sections)
create_template("Creative Design", "creative", creative_sections)
